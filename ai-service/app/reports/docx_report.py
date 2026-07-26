"""
DOCX report via python-docx. Same content as the PDF version — see
data.py — so a reader gets the same numbers regardless of which format
they open, just formatted differently.
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SEVERITY_COLORS = {
    "CRITICAL": RGBColor(0xEF, 0x44, 0x44),
    "MAJOR": RGBColor(0xF5, 0x9E, 0x0B),
    "MINOR": RGBColor(0xEA, 0xB3, 0x08),
}


def _add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def generate_docx(data: dict, output_path: str):
    doc = Document()
    project = data["project"]

    title = doc.add_heading("AI RERA Auditor", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("From Blueprint to Approval — Powered by Artificial Intelligence.")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _add_heading(doc, f"Compliance Audit Report: {project.name}", level=1)
    doc.add_paragraph(f"Location: {project.location or 'Not specified'}")
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")

    # --- Compliance summary ---
    _add_heading(doc, "Compliance Summary", level=2)
    if data["compliance_score"] is not None:
        p = doc.add_paragraph()
        run = p.add_run(f"Compliance Score: {data['compliance_score']}/100")
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph(f"Approval Probability: {data['approval_probability']}%")
    else:
        doc.add_paragraph("No compliance analysis has been run yet for this project.")

    counts = {k: len(v) for k, v in data["violations_by_severity"].items()}
    table = doc.add_table(rows=2, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Critical", "Major", "Minor", "Total"
    vals = table.rows[1].cells
    vals[0].text = str(counts["CRITICAL"])
    vals[1].text = str(counts["MAJOR"])
    vals[2].text = str(counts["MINOR"])
    vals[3].text = str(sum(counts.values()))

    # --- Building overview ---
    if data["buildings"]:
        doc.add_paragraph()
        _add_heading(doc, "Building Overview", level=2)
        b_table = doc.add_table(rows=1, cols=5)
        b_table.style = "Light Grid Accent 1"
        hdr = b_table.rows[0].cells
        for i, label in enumerate(["Name", "Type", "Floors", "FAR", "Ground Coverage %"]):
            hdr[i].text = label
        for b in data["buildings"]:
            row = b_table.add_row().cells
            row[0].text = b.name
            row[1].text = b.building_type
            row[2].text = str(b.num_floors)
            row[3].text = str(b.far_calculated) if b.far_calculated else "—"
            row[4].text = str(b.ground_coverage_pct) if b.ground_coverage_pct else "—"

    # --- Violations ---
    doc.add_paragraph()
    _add_heading(doc, "Violations", level=2)
    if not data["violations"]:
        doc.add_paragraph("No open violations.")
    else:
        for severity in ["CRITICAL", "MAJOR", "MINOR"]:
            items = data["violations_by_severity"].get(severity, [])
            if not items:
                continue
            p = doc.add_paragraph()
            run = p.add_run(f"{severity.title()} ({len(items)})")
            run.bold = True
            run.font.color.rgb = SEVERITY_COLORS[severity]

            v_table = doc.add_table(rows=1, cols=4)
            v_table.style = "Light Grid Accent 1"
            hdr = v_table.rows[0].cells
            for i, label in enumerate(["Rule", "Description", "Detected", "Required"]):
                hdr[i].text = label
            for v in items:
                row = v_table.add_row().cells
                row[0].text = v.rule.rule_code if v.rule else "—"
                row[1].text = v.description
                row[2].text = str(v.detected_value) if v.detected_value is not None else "—"
                row[3].text = str(v.required_value) if v.required_value is not None else "—"
            doc.add_paragraph()

    # --- AI Suggestions ---
    if data["suggestions"]:
        doc.add_page_break()
        _add_heading(doc, "AI Suggestions", level=2)
        for s in data["suggestions"]:
            doc.add_paragraph(s.suggestion_text, style="List Bullet")

    # --- GIS / Fire Access ---
    site = data["site"]
    if site:
        doc.add_paragraph()
        _add_heading(doc, "Site & Fire Access", level=2)
        doc.add_paragraph(f"Address: {site.geocoded_address or '—'}")
        if site.nearest_road_distance_m is not None:
            estimated_note = " (estimated)" if site.nearest_road_width_is_estimated else " (measured)"
            doc.add_paragraph(
                f"Nearest road: {site.nearest_road_name or 'Unnamed'} — "
                f"{site.nearest_road_distance_m}m away, {site.nearest_road_width_m}m wide{estimated_note}"
            )
            doc.add_paragraph(
                f"Fire tender access: {'Compliant' if site.fire_access_compliant else 'Non-compliant'}"
            )
        if site.encroachment_status == "NOT_AVAILABLE":
            doc.add_paragraph("Encroachment check: not available (requires an authoritative cadastral/survey boundary).")

    # --- Disclaimer ---
    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "This report is generated by an AI-assisted compliance tool. Rule thresholds are "
        "representative and must be verified against the current local development authority's "
        "bylaws before submission. This is not a substitute for a licensed architect's or "
        "structural engineer's certification."
    )
    disclaimer.runs[0].font.size = Pt(8)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(output_path)
